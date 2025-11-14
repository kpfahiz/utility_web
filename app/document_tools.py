import os

from flask import current_app
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from .utils.logger import logger
from .utils.timing import measure_duration

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False


@measure_duration
def convert_pdf_to_doc(file: FileStorage) -> dict:
    """
    Convert a PDF file to DOCX format.

    Args:
        file: Uploaded PDF file.

    Returns:
        Dictionary containing:
            - filename: Converted file name
            - path: Absolute path to the converted DOCX file
            - original_size: Original file size in bytes
            - converted_size: Converted file size in bytes
    """
    if not PDF2DOCX_AVAILABLE:
        raise ImportError(
            "pdf2docx library is not installed. Please install it using: pip install pdf2docx"
        )

    upload_folder = current_app.config["UPLOAD_FOLDER"]
    os.makedirs(upload_folder, exist_ok=True)

    filename = secure_filename(file.filename)
    original_path = os.path.join(upload_folder, filename)

    # Get base filename without extension
    base_name = os.path.splitext(filename)[0]
    converted_filename = f"converted_{base_name}.docx"
    converted_path = os.path.join(upload_folder, converted_filename)

    try:
        file.save(original_path)
        original_size = os.path.getsize(original_path)

        # Convert PDF to DOCX
        cv = Converter(original_path)
        cv.convert(converted_path, start=0, end=None)
        cv.close()

        converted_size = os.path.getsize(converted_path)

        return {
            "filename": converted_filename,
            "path": converted_path,
            "original_size": original_size,
            "converted_size": converted_size,
        }

    except Exception as exc:
        logger.error(f"Failed to convert PDF {filename} to DOCX: {exc}", exc_info=True)
        raise


@measure_duration
def convert_doc_to_pdf(file: FileStorage) -> dict:
    """
    Convert a DOCX file to PDF format.

    Args:
        file: Uploaded DOCX file.

    Returns:
        Dictionary containing:
            - filename: Converted file name
            - path: Absolute path to the converted PDF file
            - original_size: Original file size in bytes
            - converted_size: Converted file size in bytes
    """
    if not DOCX2PDF_AVAILABLE:
        raise ImportError(
            "Required libraries are not installed. Please install using: "
            "pip install python-docx reportlab"
        )

    upload_folder = current_app.config["UPLOAD_FOLDER"]
    os.makedirs(upload_folder, exist_ok=True)

    filename = secure_filename(file.filename)
    original_path = os.path.join(upload_folder, filename)

    # Get base filename without extension
    base_name = os.path.splitext(filename)[0]
    converted_filename = f"converted_{base_name}.pdf"
    converted_path = os.path.join(upload_folder, converted_filename)

    try:
        file.save(original_path)
        original_size = os.path.getsize(original_path)

        # Read DOCX file
        doc = Document(original_path)

        # Create PDF document
        pdf_doc = SimpleDocTemplate(
            converted_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # Build PDF content
        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=16,
            spaceAfter=12,
            alignment=TA_CENTER,
        )

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                story.append(Spacer(1, 0.2 * inch))
                continue

            # Determine alignment
            alignment = TA_LEFT
            if paragraph.alignment == 1:  # CENTER
                alignment = TA_CENTER
            elif paragraph.alignment == 2:  # RIGHT
                alignment = TA_RIGHT
            elif paragraph.alignment == 3:  # JUSTIFY
                alignment = TA_JUSTIFY

            # Determine style based on paragraph style
            style = styles["Normal"]
            if "Heading 1" in paragraph.style.name:
                style = title_style
            elif "Heading 2" in paragraph.style.name:
                style = styles["Heading2"]
            elif "Heading 3" in paragraph.style.name:
                style = styles["Heading3"]

            # Create paragraph for PDF
            para = Paragraph(text, style)
            para.alignment = alignment
            story.append(para)
            story.append(Spacer(1, 0.1 * inch))

        # Handle tables
        for table in doc.tables:
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors

            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)

            if table_data:
                pdf_table = Table(table_data)
                pdf_table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE", (0, 0), (-1, 0), 12),
                            ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                            ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                            ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ]
                    )
                )
                story.append(Spacer(1, 0.2 * inch))
                story.append(pdf_table)
                story.append(Spacer(1, 0.2 * inch))

        # Build PDF
        pdf_doc.build(story)

        converted_size = os.path.getsize(converted_path)

        return {
            "filename": converted_filename,
            "path": converted_path,
            "original_size": original_size,
            "converted_size": converted_size,
        }

    except Exception as exc:
        logger.error(f"Failed to convert DOCX {filename} to PDF: {exc}", exc_info=True)
        raise

